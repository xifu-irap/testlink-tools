#!/usr/bin/env python
# -*- coding: utf-8 -*-
#
#     Copyright (c) Nolhan Dumoulin, IRAP Toulouse
#
#     This program is free software: you can redistribute it and/or modify
#     it under the terms of the GNU General Public License as published by
#     the Free Software Foundation, either version 3 of the License, or
#     (at your option) any later version.
#
#     This program is distributed in the hope that it will be useful,
#     but WITHOUT ANY WARRANTY; without even the implied warranty of
#     MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#     GNU General Public License for more details.
#
#     You should have received a copy of the GNU General Public License
#     along with this program.  If not, see <https://www.gnu.org/licenses/>.
#
#     testcasev1.py
#

from docx.opc.exceptions import PackageNotFoundError
from docx.api import Document # function to open a docx file
from docx.document import Document as _Document # Document object (created with Document function)
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph

"""This module manages the X-IFU/DRE test procedures documents

1. Open test.docx and read paragraphs and tables
2. Assert that :
    the doc is organized into chapters, one chapter (style = "Heading 1") contains the test procedure
    each sub-chapter (style = "Heading 2") corresponds to a "Test Suite" for TestLink
    every "Test Suite" contains one or more "Test Case"
3. Convert from WORD to XML for TestLink ==> docx_to_XML (DONE)
"""

class TestSuite:
    """ Manage Test Suite sections of the test procedure """

    XML_TS_START = '<testsuite name="{}"> \n'
    XML_TS_STOP = '</testsuite>'

    def __init__(self,name,details,num_ts,tc_list):
        """ Constructor """
        self.name = name
        self.details = details
        self.order = num_ts
        self.tc_list = tc_list
        self.xml_text =""

    def to_xml(self):
        """ generates the equivalent xml text of the Test Suite """ 

        self.xml_text += TestSuite.XML_TS_START.format(self.name)
        self.xml_text += "<node_order><![CDATA[{}]]></node_order> \n".format(self.order)
        self.xml_text += "<details><![CDATA[{}]]></details> \n".format(self.details)
        
        # Generates the xml text of the Test Cases in the Test Suite"
        for testcase in self.tc_list:
            self.xml_text += testcase.to_xml()
        self.xml_text += TestSuite.XML_TS_STOP
        return self.xml_text


class TestCase :
    """ Manage Test Cases of the test procedure """

    XML_TC_START = '<testcase name="{}"> \n'
    XML_TC_STOP = "</testcase>"
    XML_STEP_START = "<step> \n"
    XML_STEP_STOP = "</step> \n"

    def __init__(self,table,num_tc):
        """ Constructor """

        self.table = table
        self.xml_text = ""
        self.xml_tags = ['node_order','preconditions','steps']
        self.xml_vals = [num_tc,"",""]
        self.xml_dict = dict(zip(self.xml_tags,self.xml_vals))
        
    def to_xml(self):
        """ generates the equivalent xml text of the Test Case """

        nb_row = 0
        step_counter = 0

        # read the table row by row
        for row in self.table.rows:
            
            nb_cell = 0
            nb_row += 1

            # row 2 and 4 are ignored
            if nb_row == 2 or nb_row == 4: 
                continue
            # the last two rows are ignored
            if nb_row == len(self.table.rows) - 1:
                break

            # read the cells of each row
            for cell in row.cells:
                nb_cell +=1
                cell_txt =""
                # pickup the name
                if nb_row == 1:
                    self.xml_text += TestCase.XML_TC_START.format(cell.text)
                    break
                # pickup preconditions
                elif nb_row == 3:
                    # manage the line breaks
                    for para in cell.paragraphs:
                        if para.text != "":
                            cell_txt += '<p>' + para.text + '</p>' +'\n'
                    self.xml_dict["preconditions"] = "<![CDATA[{}]]>".format(cell_txt)
                    break
                # read the steps section of the Test case
                else:
                    # add the step number
                    if nb_cell == 1 :
                        self.xml_dict["steps"] += TestCase.XML_STEP_START
                        step_counter += 1
                        self.xml_dict["steps"] += "<{0}> {1} </{0}> \n".format("step_number",step_counter)
                    # pickup actions
                    elif nb_cell == 2 :
                        # manage the line breaks
                        for para in cell.paragraphs:
                            if para.text != "":
                                cell_txt +='<p>' + para.text + '</p>' +'\n'
                        self.xml_dict["steps"] += "<{0}> <![CDATA[{1}]]> </{0}> \n".format("actions",cell_txt + '\n')
                    # pickup expected results
                    elif nb_cell == 3:
                        # manage the line breaks
                        for para in cell.paragraphs:
                            if para.text != "":
                                cell_txt += '<p>' + para.text + '</p>' +'\n'
                        self.xml_dict["steps"] += "<{0}> <![CDATA[{1}]]> </{0}> \n".format("expectedresults",cell_txt + '\n')
                        self.xml_dict["steps"] += TestCase.XML_STEP_STOP
                        # the other cells of a step row are ignored
                        break    
        
        # Add a general step to the Test Case (not in the docx file)
        self.xml_dict["steps"] += TestCase.XML_STEP_START
        self.xml_dict["steps"] += "<{0}> {1} </{0}> \n".format("step_number",step_counter+1)
        self.xml_dict["steps"] += "<{0}> <![CDATA[{1}]]> </{0}> \n".format("actions", "Lister les participants au test" + '\n')
        self.xml_dict["steps"] += "<{0}> <![CDATA[{1}]]> </{0}> \n".format("expectedresults", "" + '\n')
        self.xml_dict["steps"] += TestCase.XML_STEP_STOP

        # generates the xml text of the Test Case
        for tag,val in self.xml_dict.items():
            self.xml_text +=  "<{0}> {1} </{0}> \n".format(tag,val)
        self.xml_text += TestCase.XML_TC_STOP
        return self.xml_text

                           
class DocXML:
    """ Manage Word (docx) test procedure file """ 

    XML_DOC_START = """<?xml version="1.0" encoding="UTF-8"?>
    <testsuite id="" name="" >
    <node_order><![CDATA[]]>
    </node_order><details>
    <![CDATA[]]></details>"""
    XML_DOC_STOP = "</testsuite>"

    def __init__(self,filename):
        """ Constructor """

        try:
            self.doc = Document(filename)
            print("document", self.doc)
            
        except (ValueError, PackageNotFoundError) as error:
            print("ERROR {0} : {1}".format(type(error), error))

        self.xml_text = DocXML.XML_DOC_START
        self.ts_list = []

    def __iter_block_items(self, parent):
        """From : https://stackoverflow.com/questions/29240707/python-docx-get-tables-from-paragraph
        
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        elif isinstance(parent, _Row):
            parent_elm = parent._tr
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def docx_to_xml(self):
        """ Scan a docx document to extract Paragraphs and Tables for "Test Suites" and "Test Cases" and generates the full xml text"""

        new_ts = False
        details_ts =""
        name_ts= ""
        tc_list = []
        TS_counter = 0
        TC_counter = 0

        # one pass to read the word document and pickup the Test Suites and Test Cases informations
        for block in self.__iter_block_items(self.doc):
            # read paragraphs
            if isinstance(block, Paragraph):
                # find a new Test Suite and pickup the title as Test Suite name
                if block.style.name.lower().find("heading 2") >=0 and block.text.lower().find("test suite") >= 0 :

                    # if a new Test Suite is find, add the previous one to the list and reset the informations
                    if new_ts == False and TS_counter > 0:
                        self.ts_list.append(TestSuite(name_ts, details_ts,TS_counter,tc_list))
                        name_ts = ""
                        details_ts = ""
                        tc_list = []
                    
                    # pickup the title
                    new_ts = True
                    TS_counter +=1
                    name_ts = block.text
                
                # if a Test Suite is open pickup the following text as details 
                elif new_ts and block.text !="":
                    details_ts += '<p>' + block.text + '</p>' +'\n'
            # read table
            elif isinstance(block, Table):
                # find a new Test case and pickup the corresponding table 
                if block.cell(0,0).text.lower().find("test case") >= 0 and block.cell(0,0).paragraphs[0].style.name.lower().find("heading 3") >= 0:
                    
                    # details of the Test Suite is frozen and pickup the table
                    new_ts = False
                    TC_counter += 1
                    tc_list.append(TestCase(block,TC_counter))

        # add the last Test Suite to the list
        self.ts_list.append(TestSuite(name_ts, details_ts,TS_counter,tc_list))
    
        # generates the full xml text
        for testsuite in self.ts_list:
            self.xml_text += testsuite.to_xml()
        self.xml_text += DocXML.XML_DOC_STOP
        return self.xml_text    
    

if __name__ == "__main__":
    
    import easygui

    # Select input test procedure (docx file)
    filename =  easygui.fileopenbox("Please select a file",default="C:")

    #convert to XML text for TestLink
    try:
        document = DocXML(filename)
        xml_text = document.docx_to_xml() 
        f = open("result.xml", 'w', encoding='utf-8')
        f.write(xml_text)
        f.close()
    except:
        print("Please choose a valid .docx file...")
else:
    print("Importing...", __name__)





